import io
import os
import tempfile
import zipfile

import pytest
from fpdf import FPDF
from docx import Document as DocxDocument
from app.ingestion.parsers import _parse_docx, parse_file, Page, OcrRequiredError, UnsupportedFileError


def _write(path, content):
    with open(path, "w") as f:
        f.write(content)


def test_html_parser_extracts_text_and_pages():
    with tempfile.TemporaryDirectory() as d:
        p = os.path.join(d, "manual.html")
        _write(p, "<html><body><p>Hello world</p></body></html>")
        pages = parse_file(p, "manual.html")
    assert pages == [Page(number=1, text="Hello world")]


def test_docx_parser_extracts_text():
    with tempfile.TemporaryDirectory() as d:
        p = os.path.join(d, "m.docx")
        doc = DocxDocument()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(p)
        pages = parse_file(p, "m.docx")
    assert len(pages) == 1
    assert "First paragraph" in pages[0].text
    assert "Second paragraph" in pages[0].text


def test_pdf_parser_extracts_text_per_page():
    with tempfile.TemporaryDirectory() as d:
        p = os.path.join(d, "m.pdf")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        pdf.cell(0, 10, "Page one text")
        pdf.add_page()
        pdf.cell(0, 10, "Page two text")
        pdf.output(p)
        pages = parse_file(p, "m.pdf")
    assert [pg.number for pg in pages] == [1, 2]
    assert "Page one text" in pages[0].text
    assert "Page two text" in pages[1].text


def test_pdf_with_no_text_raises_ocr_required():
    # A PDF whose page has no extractable text simulates a scanned PDF.
    with tempfile.TemporaryDirectory() as d:
        p = os.path.join(d, "m.pdf")
        pdf = FPDF()
        pdf.add_page()
        pdf.output(p)  # blank page, no text
        with pytest.raises(OcrRequiredError):
            parse_file(p, "m.pdf")


def test_unsupported_extension_raises():
    with tempfile.TemporaryDirectory() as d:
        p = os.path.join(d, "m.txt")
        _write(p, "hello")
        with pytest.raises(UnsupportedFileError):
            parse_file(p, "m.txt")


def _make_minimal_docx(path, content):
    """Build a minimal valid .docx (OPC zip) with the given word/document.xml content."""
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '</Types>'
        ))
        z.writestr("_rels/.rels", (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            '</Relationships>'
        ))
        z.writestr("word/_rels/document.xml.rels", (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
        ))
        z.writestr("word/document.xml", content)


def test_docx_zip_bomb_rejected(tmp_path, monkeypatch):
    from app.config import get_settings
    monkeypatch.setenv("DOCX_MAX_DECOMPRESSED_BYTES", "1024")
    get_settings.cache_clear()
    # a real .docx is a zip; build a tiny one with one huge member
    bomb_path = tmp_path / "bomb.docx"
    _make_minimal_docx(bomb_path, "A" * (10 * 1024))
    with pytest.raises(UnsupportedFileError):
        _parse_docx(str(bomb_path))
